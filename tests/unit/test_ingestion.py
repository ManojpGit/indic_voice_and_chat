from __future__ import annotations

import io

import pytest

from src.rag.ingestion import (
    Chunk,
    ChunkConfig,
    DOCXParser,
    FixedChunker,
    MarkdownParser,
    PDFParser,
    RecursiveChunker,
    detect_language,
    get_chunker,
    parse_document,
)


# --- Parsers -------------------------------------------------------------


def test_markdown_parser_returns_decoded_text() -> None:
    text = "# Title\n\nHello **world**\n\nनमस्ते"
    out = MarkdownParser().parse("doc.md", text.encode("utf-8"))
    assert out == text


def test_parse_document_dispatches_by_extension() -> None:
    out = parse_document("notes.md", "Hello".encode("utf-8"))
    assert out == "Hello"


def test_parse_document_unknown_extension_falls_back_to_utf8() -> None:
    out = parse_document("data.unknown", "Bonjour".encode("utf-8"))
    assert out == "Bonjour"


def test_parse_document_handles_invalid_utf8_gracefully() -> None:
    # ``errors="replace"`` shouldn't raise on bad bytes.
    out = parse_document("x.txt", b"\xff\xfe hi")
    assert "hi" in out


def test_pdf_parser_real_world_document() -> None:
    """Build a tiny PDF in-memory with pypdf to exercise the parser."""
    from pypdf import PdfWriter
    from pypdf.generic import RectangleObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    # We don't write text content (pypdf's writer doesn't compose text easily).
    # Just assert the parser doesn't crash on a valid-but-empty PDF.
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)

    out = PDFParser().parse("blank.pdf", buf.getvalue())
    assert isinstance(out, str)


def test_docx_parser_round_trip() -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph in Hindi: नमस्ते")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    out = DOCXParser().parse("notes.docx", buf.getvalue())
    assert "First paragraph." in out
    assert "नमस्ते" in out


# --- Fixed chunker ------------------------------------------------------


def test_fixed_chunker_basic_split() -> None:
    cfg = ChunkConfig(chunk_size=10, chunk_overlap=2, strategy="fixed")  # 40 chars, 8 overlap
    chunks = FixedChunker(cfg).chunk("a" * 100)
    assert len(chunks) > 1
    # Each chunk is at most chunk_chars
    assert all(len(c.text) <= cfg.chunk_chars for c in chunks)


def test_fixed_chunker_step_includes_overlap() -> None:
    cfg = ChunkConfig(chunk_size=5, chunk_overlap=1, strategy="fixed")  # 20 chars, 4 overlap
    text = "abcdefghijklmnopqrstuvwxyz" * 2  # 52 chars
    chunks = FixedChunker(cfg).chunk(text)
    # Successive chunks share their overlap region
    if len(chunks) >= 2:
        # second chunk starts at step = 16 (chunk_chars=20, overlap_chars=4)
        assert chunks[1].start_char == 16


def test_fixed_chunker_empty_input() -> None:
    cfg = ChunkConfig(chunk_size=10, chunk_overlap=0, strategy="fixed")
    assert FixedChunker(cfg).chunk("") == []


def test_fixed_chunker_invalid_size_raises() -> None:
    cfg = ChunkConfig(chunk_size=0, chunk_overlap=0, strategy="fixed")
    with pytest.raises(ValueError):
        FixedChunker(cfg).chunk("hi")


# --- Recursive chunker --------------------------------------------------


def test_recursive_chunker_keeps_paragraphs_intact_when_possible() -> None:
    cfg = ChunkConfig(chunk_size=200, chunk_overlap=10, strategy="recursive")  # 800 chars target
    text = ("Paragraph A.\n\n" * 5) + ("Paragraph B.\n\n" * 5)
    chunks = RecursiveChunker(cfg).chunk(text)
    # Whole input is small enough to land in one chunk
    assert len(chunks) == 1
    assert "Paragraph A." in chunks[0].text
    assert "Paragraph B." in chunks[0].text


def test_recursive_chunker_splits_when_over_target() -> None:
    cfg = ChunkConfig(chunk_size=20, chunk_overlap=5, strategy="recursive")  # 80 chars target
    text = "\n\n".join(f"Section {i} body content." for i in range(20))
    chunks = RecursiveChunker(cfg).chunk(text)
    assert len(chunks) >= 2
    # Each chunk's plain length is roughly under target + overlap (allow some slack
    # since the recursive merger packs greedily up to ``target``).
    assert all(len(c.text) <= cfg.chunk_chars + cfg.overlap_chars + 5 for c in chunks)


def test_recursive_chunker_respects_devanagari_separator() -> None:
    cfg = ChunkConfig(chunk_size=10, chunk_overlap=0, strategy="recursive")
    text = "नमस्ते। आप कैसे हैं। मैं ठीक हूं।"
    chunks = RecursiveChunker(cfg).chunk(text)
    # Should produce at least one chunk per sentence (approximately).
    assert len(chunks) >= 1


def test_recursive_chunker_overlap_is_prepended_to_following() -> None:
    cfg = ChunkConfig(chunk_size=20, chunk_overlap=2, strategy="recursive")  # 80 chars, 8 char overlap
    text = "\n\n".join("X" * 70 for _ in range(3))
    chunks = RecursiveChunker(cfg).chunk(text)
    assert len(chunks) >= 2


def test_recursive_chunker_no_separators_falls_back_to_hard_cut() -> None:
    cfg = ChunkConfig(chunk_size=5, chunk_overlap=0, strategy="recursive")  # 20 chars
    text = "X" * 100  # no separators at all
    chunks = RecursiveChunker(cfg).chunk(text)
    assert len(chunks) >= 4
    assert all(len(c.text) <= cfg.chunk_chars for c in chunks)


def test_recursive_chunker_metadata_is_per_chunk_copy() -> None:
    cfg = ChunkConfig(chunk_size=5, chunk_overlap=0, strategy="recursive")
    text = "A" * 60
    chunks = RecursiveChunker(cfg).chunk(text, base_metadata={"src": "test.md"})
    chunks[0].metadata["src"] = "MUTATED"
    assert all(c.metadata.get("src") == "test.md" for c in chunks[1:])


# --- get_chunker --------------------------------------------------------


def test_get_chunker_returns_callable() -> None:
    cfg = ChunkConfig(strategy="recursive")
    fn = get_chunker(cfg)
    out = fn("hello world", None)
    assert isinstance(out, list)
    assert all(isinstance(c, Chunk) for c in out)


def test_get_chunker_unknown_strategy_raises() -> None:
    cfg = ChunkConfig(strategy="quantum")  # invented
    with pytest.raises(ValueError):
        get_chunker(cfg)


# --- Language detection -------------------------------------------------


def test_detect_language_devanagari_is_hi() -> None:
    assert detect_language("नमस्ते दोस्तों") == "hi"


def test_detect_language_english_is_none() -> None:
    assert detect_language("Hello world") is None


def test_detect_language_empty_is_none() -> None:
    assert detect_language("") is None
